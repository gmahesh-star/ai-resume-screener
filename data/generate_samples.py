from docx import Document

def create_resume(filename, name, summary, skills, experience):
    doc = Document()
    doc.add_heading(f"{name}'s Resume", 0)
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph(", ".join(skills))
    
    doc.add_heading('Experience', level=1)
    doc.add_paragraph(experience)
    
    doc.save(filename)
    print(f"Created {filename}")

if __name__ == "__main__":
    # Resume 1: Good Match
    create_resume(
        "data/Jane_Doe_Resume.docx",
        "Jane Doe",
        "Senior AI Engineer with 5 years of experience building scalable ML solutions.",
        ["Python", "Machine Learning", "NLP", "FastAPI", "spaCy", "Scikit-learn", "Git", "Docker", "SQL"],
        "Built REST APIs using FastAPI for ML model serving. Used spaCy and scikit-learn for text classification and NLP tasks. Managed CI/CD pipelines with Docker."
    )
    
    # Resume 2: Partial Match
    create_resume(
        "data/John_Smith_Resume.docx",
        "John Smith",
        "Full Stack Developer transitioning to Data Science.",
        ["Python", "JavaScript", "React", "HTML", "CSS", "MySQL", "Pandas"],
        "Built web applications using React and Python backend. Exploring machine learning techniques."
    )

    # Resume 3: Poor Match
    create_resume(
        "data/Alice_Jones_Resume.docx",
        "Alice Jones",
        "Frontend Engineer specialized in mobile and web experiences.",
        ["JavaScript", "React Native", "Swift", "Kotlin", "Figma"],
        "Developed mobile applications using React Native. Designed UI mockups using Figma."
    )
